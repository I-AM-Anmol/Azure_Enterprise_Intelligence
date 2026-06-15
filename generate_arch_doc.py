from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colour palette
NAVY      = RGBColor(0x1E, 0x3A, 0x8A)
BLUE      = RGBColor(0x25, 0x63, 0xEB)
DARK_GREY = RGBColor(0x0F, 0x17, 0x2A)
MID_GREY  = RGBColor(0x47, 0x55, 0x69)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GREEN     = RGBColor(0x06, 0x60, 0x27)

# ── Helpers
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="D1D5DB"):
    tc      = cell._tc
    tcPr    = tc.get_or_add_tcPr()
    tcBords = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBords.append(b)
    tcPr.append(tcBords)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "24")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "2563EB")
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    return p

def body(text, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.size      = Pt(10)
    run.font.color.rgb = DARK_GREY
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent  = Inches(0.25)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold           = True
        rb.font.size      = Pt(10)
        rb.font.color.rgb = DARK_GREY
        rt = p.add_run(text)
        rt.font.size      = Pt(10)
        rt.font.color.rgb = MID_GREY
    else:
        run = p.add_run(text)
        run.font.size      = Pt(10)
        run.font.color.rgb = DARK_GREY
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.25)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "1E3A8A")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name      = "Courier New"
    run.font.size      = Pt(8)
    run.font.color.rgb = WHITE
    return p

def make_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style         = "Table Grid"
    t.alignment     = WD_TABLE_ALIGNMENT.LEFT
    t.allow_autofit = False
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, "1E3A8A")
        set_cell_borders(cell, "2563EB")
        p   = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(h)
        run.bold           = True
        run.font.size      = Pt(9)
        run.font.color.rgb = WHITE
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row_data in enumerate(rows):
        drow = t.rows[ri + 1]
        bg   = "EFF6FF" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = drow.cells[ci]
            set_cell_bg(cell, bg)
            set_cell_borders(cell, "BFDBFE")
            p   = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            run = p.add_run(val)
            run.font.size      = Pt(9)
            run.font.color.rgb = DARK_GREY
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Inches(w)
    return t

def spacer(pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run("")
    run.font.size = Pt(pts)

# ═══════════════════════════════════════════════════════════════
#  COVER BANNER
# ═══════════════════════════════════════════════════════════════
for line_text, font_size, font_color, fill in [
    ("  CloudLens — MedInsight FinOps Intelligence", 20, WHITE,                     "1E3A8A"),
    ("  Architecture Summary  ·  Security Team Review", 11, RGBColor(0x93,0xC5,0xFD), "1E3A8A"),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill)
    pPr.append(shd)
    run = p.add_run(line_text)
    run.bold           = (font_size == 20)
    run.font.size      = Pt(font_size)
    run.font.color.rgb = font_color

spacer(10)

meta = doc.add_paragraph()
meta.paragraph_format.space_before = Pt(0)
meta.paragraph_format.space_after  = Pt(14)
for label, value in [
    ("Application",    "CloudLens — MedInsight FinOps Intelligence  "),
    ("Classification", "Internal  "),
    ("Date",           "June 2026  "),
    ("Owner",          "Azure FinOps Team"),
]:
    lb = meta.add_run(f"{label}: ")
    lb.bold           = True
    lb.font.size      = Pt(9)
    lb.font.color.rgb = MID_GREY
    vr = meta.add_run(value)
    vr.font.size      = Pt(9)
    vr.font.color.rgb = DARK_GREY

# ═══════════════════════════════════════════════════════════════
#  1. OVERVIEW
# ═══════════════════════════════════════════════════════════════
heading1("1. Overview")
body(
    "CloudLens is a four-layer Azure FinOps platform that collects cost and budget data "
    "from multiple Azure tenants, stores it in Microsoft Fabric Lakehouse (OneLake), "
    "manages all source code through a GitHub repository, and surfaces insights through "
    "a Streamlit web application deployed via CI/CD pipeline."
)
spacer(4)
body("The system spans the following three tenants:")
bullet("MedInsight Production  (tenant b2e2e6d4-...)")
bullet("MedInsight Engineering  (tenant ff6598a3-...)")
bullet("Milliman Inc.  (tenant e240d61e-...)  — configured, credentials pending")
spacer(6)

# ═══════════════════════════════════════════════════════════════
#  2. ARCHITECTURE DESIGN
# ═══════════════════════════════════════════════════════════════
heading1("2. Architecture Design")
body(
    "The architecture follows a Collect → Store → Deploy → Report pipeline. "
    "GitHub is the single source of truth for all code. "
    "No code is deployed manually — all changes flow through the CI/CD pipeline."
)
spacer(6)

heading2("End-to-End Data Flow")
for line in [
    "┌──────────────────────────────────────────────────────────────────┐",
    "│           AZURE TENANTS  (3 tenants)                             │",
    "│  MedInsight Production · MedInsight Engineering · Milliman       │",
    "│                                                                  │",
    "│  ┌────────────────────────┐   ┌──────────────────────────┐      │",
    "│  │  Cost Management API   │   │  Consumption/Budgets API │      │",
    "│  └──────────┬─────────────┘   └─────────────┬────────────┘      │",
    "└─────────────┼───────────────────────────────┼───────────────────┘",
    "              │  HTTPS REST (OAuth 2.0 SP)      │",
    "              ▼                                 ▼",
    "┌──────────────────────────────────────────────────────────────────┐",
    "│   LAYER 1 — FABRIC NOTEBOOK  (Synapse PySpark)                   │",
    "│   Auth: ClientSecretCredential (SP per tenant, auto-renew)       │",
    "│   Logic: list subs → cost → budgets → compute KPIs               │",
    "│   Out:   BudgetData.json  +  BudgetData (Delta table)            │",
    "└──────────────────────────┬───────────────────────────────────────┘",
    "                           │  OneLake write (blobfuse / SDK)",
    "                           ▼",
    "┌──────────────────────────────────────────────────────────────────┐",
    "│   LAYER 2 — FABRIC LAKEHOUSE  (OneLake)                          │",
    "│   /Tables/BudgetData   (Delta Lake — Parquet + tx log)           │",
    "│   /Files/BudgetData.json   (flat JSON snapshot)                  │",
    "└──────────────────────────────────────────────────────────────────┘",
    "",
    "┌──────────────────────────────────────────────────────────────────┐",
    "│   LAYER 3 — GITHUB REPOSITORY  (Source of Truth)                 │",
    "│                                                                  │",
    "│   /app.py                   Streamlit entry point                │",
    "│   /pages/                   Dashboard page modules               │",
    "│   /FABRIC collection Codes/ Fabric notebook scripts              │",
    "│   /.github/workflows/       CI/CD pipeline definitions           │",
    "│                                                                  │",
    "│   Branches:  main (production)  ·  dev (development)            │",
    "│   Pipeline:  push to main → GitHub Actions → deploy Streamlit   │",
    "└──────────────────────────┬───────────────────────────────────────┘",
    "                           │  GitHub Actions CI/CD",
    "                           │  (lint → test → deploy)",
    "                           ▼",
    "┌──────────────────────────────────────────────────────────────────┐",
    "│   LAYER 4 — STREAMLIT APP  (CloudLens)                           │",
    "│   Two environments:  Production  |  Development                  │",
    "│   Pages: Home / Budget Analysis / Storage Lifecycle /            │",
    "│           AI Spend Analyzer                                      │",
    "│   Auth to Fabric: MSAL → Power BI API → BudgetData Delta table  │",
    "└──────────────────────────┬───────────────────────────────────────┘",
    "                           │  HTTPS",
    "                           ▼",
    "                   [ User Browser ]",
]:
    code_block(line)

spacer(10)

# ═══════════════════════════════════════════════════════════════
#  3. LAYER 1 — DATA COLLECTION
# ═══════════════════════════════════════════════════════════════
heading1("3. Layer 1 — Data Collection  (Microsoft Fabric Notebook)")
make_table(
    ["Property", "Value"],
    [
        ["Runtime",        "Microsoft Fabric Notebook — Synapse PySpark kernel"],
        ["Lakehouse",      "MedInsight_Azure_FinOps_Intelligence"],
        ["Workspace ID",   "eca3c81e-a968-42a5-899f-d8fc1a45ebec"],
        ["Lakehouse ID",   "84677723-0231-4f98-b71a-14879da990b8"],
        ["Trigger",        "Manual on-demand refresh (pipeline/scheduled trigger planned)"],
        ["SDK",            "azure-identity · azure-mgmt-resource · requests"],
        ["Code location",  "GitHub: /FABRIC collection Codes/FABRIC_BudgetAnalyzer.py"],
    ],
    col_widths=[2.0, 4.5]
)
spacer(8)

heading2("Processing Steps")
for prefix, detail in [
    ("1. Authenticate",      "Instantiates ClientSecretCredential per tenant SP. Tokens are callable (not static strings) — auto-refresh after 1-hour expiry prevents HTTP 401 on long-running jobs."),
    ("2. List Subscriptions","Calls management.azure.com/subscriptions to enumerate all active subscriptions per tenant."),
    ("3. Query Cost API",    "Posts to Microsoft.CostManagement/query per subscription. Retrieves month-to-date actual spend in USD. Sequential with 6-second gap to stay under ~10 req/min SP rate limit."),
    ("4. Query Budgets API", "Calls Microsoft.Consumption/budgets per subscription. Retrieves thresholds, alert levels, and contact addresses. Up to 50 parallel workers (no rate limit on this API)."),
    ("5. Compute KPIs",      "Calculates: budget utilisation %, days until alert breach, days remaining in billing month, over-budget flag."),
    ("6. Write to Lakehouse","Writes BudgetData Delta table (overwriteSchema=true) and BudgetData.json to OneLake. Falls back to notebookutils.fs.put or OneLake REST API if blobfuse mount is unavailable."),
]:
    bullet(f"  {detail}", bold_prefix=prefix)

spacer(8)

heading2("Azure APIs Called")
make_table(
    ["API Endpoint", "Scope", "Purpose"],
    [
        ["management.azure.com/subscriptions",                       "Per tenant",       "List all subscriptions"],
        ["management.azure.com/.../Microsoft.CostManagement/query",  "Per subscription", "Month-to-date actual spend (USD)"],
        ["management.azure.com/.../Microsoft.Consumption/budgets",   "Per subscription", "Budget limits and alert thresholds"],
    ],
    col_widths=[3.0, 1.5, 2.0]
)
spacer(8)

heading2("Tenant & Service Principal Configuration")
make_table(
    ["Tenant", "Tenant ID (prefix)", "Auth Method", "SP Client ID (prefix)"],
    [
        ["MedInsight Production",  "b2e2e6d4-...", "Service Principal (Client Credentials)", "37b785ce-..."],
        ["MedInsight Engineering", "ff6598a3-...", "Service Principal (Client Credentials)", "c498ca52-..."],
        ["Milliman Inc.",          "e240d61e-...", "Not yet configured",                     "—"],
    ],
    col_widths=[1.8, 1.5, 2.5, 1.7]
)
spacer(8)

heading2("Rate Limiting & Retry Logic")
bullet("Cost Management API: ~10 requests/minute per Service Principal token.")
bullet("Strategy: sequential queries with 6-second inter-request sleep.")
bullet("HTTP 429: reads Retry-After header; backs off and retries up to 20 attempts.")
bullet("HTTP 401: callable token pattern forces fresh token acquisition on every request.")
spacer(6)

# ═══════════════════════════════════════════════════════════════
#  4. LAYER 2 — STORAGE
# ═══════════════════════════════════════════════════════════════
heading1("4. Layer 2 — Storage  (Microsoft Fabric / OneLake)")
make_table(
    ["Component", "Detail"],
    [
        ["Platform",         "Microsoft Fabric (OneLake)"],
        ["Lakehouse Name",   "MedInsight_Azure_FinOps_Intelligence"],
        ["Lakehouse ID",     "84677723-0231-4f98-b71a-14879da990b8"],
        ["Workspace ID",     "eca3c81e-a968-42a5-899f-d8fc1a45ebec"],
        ["Primary format",   "Delta Lake (Parquet + transaction log) — /Tables/BudgetData"],
        ["Secondary format", "Flat JSON snapshot — /Files/BudgetData.json"],
        ["Write mode",       "Full overwrite on each refresh (rolling 7-day upsert planned)"],
        ["Schema evolution", "overwriteSchema=true — allows column type changes between runs"],
        ["Access control",   "Microsoft Fabric Workspace RBAC"],
        ["Data sensitivity", "Subscription names, USD cost figures, budget thresholds — no PII"],
    ],
    col_widths=[2.0, 4.5]
)
spacer(6)

# ═══════════════════════════════════════════════════════════════
#  5. LAYER 3 — GITHUB REPOSITORY & CI/CD
# ═══════════════════════════════════════════════════════════════
heading1("5. Layer 3 — GitHub Repository & CI/CD Pipeline")
body(
    "All source code — Streamlit application pages, Fabric notebook scripts, and pipeline "
    "definitions — is stored in the MedInsight GitHub repository. "
    "No code is deployed manually. Every change to production flows through a GitHub Actions "
    "workflow that lints, tests, and deploys the application automatically."
)
spacer(6)

heading2("Repository Structure")
make_table(
    ["Path", "Contents"],
    [
        ["app.py",                           "Streamlit application entry point"],
        ["pages/0_Home.py",                  "Home / navigation page"],
        ["pages/1_Storage_Lifecycle.py",     "Storage cost analysis dashboard"],
        ["pages/2_Budget_Analysis.py",       "Budget utilisation and alert dashboard"],
        ["pages/3_AI_Spend_Analyzer.py",     "Azure AI/ML spend breakdown"],
        ["FABRIC collection Codes/",         "Fabric notebook Python scripts (BudgetAnalyzer)"],
        [".github/workflows/deploy.yml",     "GitHub Actions CI/CD workflow definition"],
        ["requirements.txt",                 "Python dependency list"],
    ],
    col_widths=[2.8, 3.7]
)
spacer(8)

heading2("Branching Strategy")
make_table(
    ["Branch", "Purpose", "Deploys To"],
    [
        ["main",      "Stable production-ready code — protected branch, requires PR review", "Production Streamlit environment"],
        ["dev",       "Active development and testing — changes reviewed here before merging to main", "Development Streamlit environment"],
        ["feature/*", "Individual feature or fix branches — merged into dev via pull request", "No automatic deployment"],
    ],
    col_widths=[1.3, 3.2, 2.0]
)
spacer(8)

heading2("CI/CD Pipeline — GitHub Actions")
body(
    "A push or merged pull request to main (or dev) triggers the following automated pipeline:"
)
spacer(4)
for line in [
    "  GitHub Push / PR merge to main",
    "         │",
    "         ▼",
    "  ┌─────────────────────────────────────┐",
    "  │  Step 1 — Lint & Static Analysis    │  flake8 / pylint",
    "  │  Step 2 — Unit Tests                │  pytest (data validation checks)",
    "  │  Step 3 — Deploy to Streamlit       │  streamlit deploy / Azure Web App",
    "  └─────────────────────────────────────┘",
    "         │",
    "         ▼",
    "  Production Streamlit App updated — no manual steps required",
]:
    code_block(line)

spacer(8)

heading2("Deployment Environments")
make_table(
    ["Environment", "Branch", "Purpose", "Access"],
    [
        ["Production", "main", "Live environment used by finance and engineering teams", "Internal users only"],
        ["Development","dev",  "Testing ground — changes validated here before production promotion", "Development team"],
    ],
    col_widths=[1.3, 1.0, 3.0, 1.6]
)
spacer(8)

heading2("GitHub → Streamlit Connection")
body(
    "The Streamlit application is connected directly to the GitHub repository. "
    "When the CI/CD pipeline deploys successfully, the live application is automatically "
    "updated with the latest code from the main branch. "
    "Secrets (SP credentials, Fabric workspace IDs) are stored as GitHub Actions Secrets "
    "and injected as environment variables at deploy time — never stored in code."
)
spacer(4)
bullet("GitHub Secrets store:  AZURE_CLIENT_SECRET, FABRIC_WORKSPACE_ID, FABRIC_DATASET_ID")
bullet("Streamlit app reads:  os.environ[] at runtime — no hardcoded credentials in source")
bullet("Pen test scope:       Both Production and Development environments before go-live")
spacer(6)

# ═══════════════════════════════════════════════════════════════
#  6. LAYER 4 — REPORTING
# ═══════════════════════════════════════════════════════════════
heading1("6. Layer 4 — Reporting  (Streamlit Web Application — CloudLens)")
make_table(
    ["Property", "Value"],
    [
        ["Framework",      "Python · Streamlit (multi-page application)"],
        ["App Name",       "CloudLens — MedInsight FinOps"],
        ["Auth to Fabric", "MSAL PublicClientApplication → Power BI API scopes → Fabric REST API"],
        ["Dataset ID",     "56e6e1c3-8b70-4c53-b288-331041ce1f3f"],
        ["Auto-refresh",   "Every 5 minutes (st_autorefresh interval = 300,000 ms)"],
        ["Code repo",      "MedInsight GitHub — main branch"],
        ["Deployment",     "Automated via GitHub Actions on merge to main"],
    ],
    col_widths=[2.0, 4.5]
)
spacer(8)

heading2("Application Pages")
make_table(
    ["Page", "Route", "Description"],
    [
        ["Home",              "pages/0_Home.py",             "Landing page — navigation cards to all dashboards"],
        ["Budget Analysis",   "pages/2_Budget_Analysis.py",  "KPI cards, utilisation bars, per-subscription spend vs. budget, alert countdown"],
        ["Storage Lifecycle", "pages/1_Storage_Lifecycle.py","Azure storage cost analysis and lifecycle policy review"],
        ["AI Spend Analyzer", "pages/3_AI_Spend_Analyzer.py","Azure AI/ML service spend breakdown"],
    ],
    col_widths=[1.6, 2.3, 2.6]
)
spacer(6)

# ═══════════════════════════════════════════════════════════════
#  7. ARCHITECTURE SUMMARY
# ═══════════════════════════════════════════════════════════════
heading1("7. Architecture Summary")
make_table(
    ["Layer", "Technology", "Role", "Auth Method"],
    [
        ["1 — Collection", "Microsoft Fabric Notebook (PySpark)",  "Pull cost & budget data from Azure APIs; compute KPIs",              "OAuth 2.0 Client Credentials (SP)"],
        ["2 — Storage",    "Microsoft Fabric Lakehouse (OneLake)", "Persist data as Delta table + JSON snapshot",                        "Fabric Managed Identity / Workspace RBAC"],
        ["3 — Repository", "GitHub (MedInsight instance)",         "Single source of truth for all code; CI/CD pipeline triggers deploy", "GitHub PAT / Actions Secrets"],
        ["4 — Reporting",  "Streamlit Python App (CloudLens)",     "Serve interactive dashboards; two environments (prod / dev)",         "MSAL interactive (Power BI API)"],
    ],
    col_widths=[1.3, 1.9, 2.6, 1.7]
)
spacer(8)

heading2("Network & Protocol Boundaries")
make_table(
    ["Flow", "Protocol", "Authentication"],
    [
        ["Fabric Notebook → Azure Cost Management API",   "HTTPS REST",              "OAuth 2.0 — SP Client Credentials"],
        ["Fabric Notebook → Azure Consumption API",       "HTTPS REST",              "OAuth 2.0 — SP Client Credentials"],
        ["Fabric Notebook → OneLake",                     "Internal blobfuse / SDK", "Fabric Managed Identity"],
        ["GitHub Actions → Streamlit Deploy",             "HTTPS",                   "GitHub Actions Secrets"],
        ["Streamlit App → Fabric REST API",               "HTTPS REST",              "OAuth 2.0 — MSAL interactive / device flow"],
        ["User → Streamlit App",                          "HTTPS",                   "Browser session (pen test required pre go-live)"],
    ],
    col_widths=[2.8, 1.7, 2.0]
)
spacer(10)

# ── Footer
fp = doc.add_paragraph()
fp.paragraph_format.space_before = Pt(16)
fp.paragraph_format.space_after  = Pt(0)
pPr = fp._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
tb = OxmlElement("w:top")
tb.set(qn("w:val"),   "single")
tb.set(qn("w:sz"),    "4")
tb.set(qn("w:space"), "6")
tb.set(qn("w:color"), "BFDBFE")
pBdr.append(tb)
pPr.append(pBdr)
fr = fp.add_run("CloudLens · MedInsight FinOps Intelligence · Architecture Summary · Internal Use Only · June 2026")
fr.font.size      = Pt(8)
fr.font.color.rgb = MID_GREY
fr.italic         = True

# ── Save
out = r"c:\Users\anmol.sharma\Desktop\Azure_Enterprise_Intelligence\CloudLens_Architecture_Summary.docx"
doc.save(out)
print(f"Saved: {out}")
